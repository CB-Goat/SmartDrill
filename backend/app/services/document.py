from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
import os

def ensure_upload_dir():
    upload_dir = "/app/uploads"
    if not os.path.exists(upload_dir):
        os.makedirs(upload_dir, exist_ok=True)
    return upload_dir

def generate_review_document(knowledge_points, user_id):
    doc = Document()
    
    doc.add_heading('复习资料', 0)
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    for i, kp in enumerate(knowledge_points, 1):
        doc.add_heading(f'{i}. {kp.title}', level=1)
        doc.add_paragraph(kp.content)
        
        if kp.exam_frequency:
            doc.add_paragraph(f'考试频率: {kp.exam_frequency}')
        if kp.exam_types:
            doc.add_paragraph(f'考试题型: {kp.exam_types}')
        
        doc.add_paragraph()
    
    upload_dir = ensure_upload_dir()
    file_path = os.path.join(upload_dir, f'review_{user_id}_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx')
    doc.save(file_path)
    
    return file_path

def generate_practice_document(questions, user_id):
    doc = Document()
    
    doc.add_heading('练习题', 0)
    doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    for i, q in enumerate(questions, 1):
        doc.add_heading(f'第{i}题', level=1)
        doc.add_paragraph(q.question)
        
        if q.question_type:
            doc.add_paragraph(f'题型: {q.question_type}')
        if q.difficulty:
            doc.add_paragraph(f'难度: {q.difficulty}')
        
        doc.add_paragraph()
    
    doc.add_page_break()
    doc.add_heading('参考答案', 0)
    
    for i, q in enumerate(questions, 1):
        doc.add_heading(f'第{i}题答案', level=1)
        doc.add_paragraph(q.answer)
        doc.add_paragraph()
    
    upload_dir = ensure_upload_dir()
    file_path = os.path.join(upload_dir, f'practice_{user_id}_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx')
    doc.save(file_path)
    
    return file_path