from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from app.database import get_db
from app.models.models import User, Grade, Subject, Semester, Unit, KnowledgePoint, ExamPoint, Order, Question, QuestionType, Difficulty
from app.utils.auth import get_current_user
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from urllib.parse import quote
import random

from app.api.knowledge_import import generate_unit_word_doc

router = APIRouter(prefix="/user", tags=["用户首页"])

GRADE_ORDER = ['一年级', '二年级', '三年级', '四年级', '五年级', '六年级', '初一', '七年级', '初二', '八年级', '初三', '九年级']

def calculate_current_grade(db, user):
    if not user.child_grade_id or not user.child_grade_set_at:
        return None, 0, '上册'
    
    set_grade = db.query(Grade).filter(Grade.id == user.child_grade_id).first()
    if not set_grade:
        return None, 0, '上册'
    
    from datetime import datetime
    days_passed = (datetime.now() - user.child_grade_set_at).days
    months_passed = days_passed / 30
    
    try:
        set_idx = GRADE_ORDER.index(set_grade.name)
    except ValueError:
        current_month = datetime.now().month
        semester = '下册' if current_month >= 2 and current_month <= 7 else '上册'
        return set_grade, months_passed % 12, semester
    
    grade_offset = int(months_passed / 12)
    current_idx = min(set_idx + grade_offset, len(GRADE_ORDER) - 1)
    
    current_grade = db.query(Grade).filter(
        Grade.version_id == set_grade.version_id,
        Grade.name == GRADE_ORDER[current_idx]
    ).first()
    
    months_in_grade = months_passed % 12 if current_idx < len(GRADE_ORDER) - 1 else 12
    
    current_month = datetime.now().month
    semester = '下册' if current_month >= 2 and current_month <= 7 else '上册'
    
    return (current_grade or set_grade), months_in_grade, semester

def calculate_current_unit_number(current_month, semester):
    if semester == '上册':
        if current_month >= 9:
            return current_month - 8
        else:
            return current_month + 4
    else:
        if current_month >= 2:
            return current_month - 1
        else:
            return 6

def set_run_shading(run, color='D9D9D9'):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    rPr.append(shd)

def format_content(doc, content):
    if not content:
        return
    
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        
        if line.startswith('【') and '】' in line:
            end_idx = line.index('】')
            title_text = line[:end_idx + 1]
            rest_text = line[end_idx + 1:].strip()
            
            p = doc.add_paragraph()
            run = p.add_run(title_text)
            run.font.bold = True
            run.font.size = Pt(12)
            
            if rest_text:
                p2 = doc.add_paragraph()
                p2.paragraph_format.first_line_indent = Pt(24)
                p2.add_run(rest_text)
            
            i += 1
        elif '：' in line or ':' in line:
            colon_char = '：' if '：' in line else ':'
            colon_idx = line.index(colon_char)
            key_text = line[:colon_idx].strip()
            value_text = line[colon_idx + 1:].strip()
            
            p = doc.add_paragraph()
            key_run = p.add_run(key_text + colon_char)
            key_run.font.bold = True
            set_run_shading(key_run)
            if value_text:
                p.add_run(value_text)
            i += 1
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            p.add_run(line)
            i += 1

@router.get("/home-data")
def get_home_data(
    grade_id: int = None,
    semester: str = None,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    from datetime import datetime
    
    current_grade, months_in_grade, current_semester = calculate_current_grade(db, user)
    
    if not current_grade:
        return {"subjects": [], "grade_name": None, "semester": None, "all_grades": [], "current_grade_id": None}
    
    if grade_id:
        selected_grade = db.query(Grade).filter(Grade.id == grade_id).first()
        if selected_grade:
            current_grade = selected_grade
        if semester:
            current_semester = semester
    
    subjects = db.query(Subject).filter(Subject.grade_id == current_grade.id).all()
    
    all_grades = db.query(Grade).filter(Grade.version_id == current_grade.version_id).all()
    grade_order_map = {name: idx for idx, name in enumerate(GRADE_ORDER)}
    all_grades = sorted(all_grades, key=lambda g: grade_order_map.get(g.name, 999))
    
    result = []
    for subject in subjects:
        semesters = db.query(Semester).filter(Semester.subject_id == subject.id).all()
        if not semesters:
            continue
        
        target_semester = None
        for sem in semesters:
            if (current_semester == '上册' and ('上' in sem.name or '一' in sem.name)) or \
               (current_semester == '下册' and ('下' in sem.name or '二' in sem.name)):
                target_semester = sem
                break
        
        if not target_semester:
            for sem in semesters:
                target_semester = sem
                break
        
        if not target_semester:
            continue
        
        all_units = db.query(Unit).filter(
            Unit.semester_id == target_semester.id
        ).order_by(Unit.unit_number).all()
        
        if not all_units:
            continue
        
        units_list = []
        for unit in all_units:
            has_knowledge = db.query(KnowledgePoint).filter(KnowledgePoint.unit_id == unit.id).first() is not None
            has_8modules = unit.unit_knowledge_json is not None and isinstance(unit.unit_knowledge_json, dict) and bool(unit.unit_knowledge_json)
            has_exam = db.query(ExamPoint).filter(ExamPoint.unit_id == unit.id).first() is not None
            question_count = db.query(Question).filter(Question.unit_id == unit.id).count()
            review_downloaded = db.query(Order).filter(
                Order.user_id == user.id,
                Order.order_type == 'knowledge',
                Order.title.like(f"%{unit.name}%")
            ).first() is not None
            practice_downloaded = db.query(Order).filter(
                Order.user_id == user.id,
                Order.order_type == 'practice',
                Order.title.like(f"%{unit.name}%")
            ).first() is not None
            units_list.append({
                "id": unit.id,
                "name": unit.name,
                "unit_number": unit.unit_number,
                "semester_name": target_semester.name,
                "has_knowledge": has_knowledge,
                "has_8modules": has_8modules,
                "has_exam": has_exam,
                "question_count": question_count,
                "review_downloaded": review_downloaded,
                "practice_downloaded": practice_downloaded
            })
        
        if units_list:
            result.append({
                "subject_id": subject.id,
                "subject_name": subject.name,
                "grade_name": current_grade.name,
                "semester": current_semester,
                "units": units_list
            })
    
    return {
        "subjects": result, 
        "grade_name": current_grade.name, 
        "semester": current_semester,
        "all_grades": [{"id": g.id, "name": g.name} for g in all_grades],
        "current_grade_id": current_grade.id
    }

@router.get("/unit-word/{unit_id}")
def get_unit_word(
    unit_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    unit = db.query(Unit).filter(Unit.id == unit_id).first()
    if not unit:
        raise HTTPException(status_code=404, detail="单元不存在")
    
    semester = db.query(Semester).filter(Semester.id == unit.semester_id).first()
    subject = db.query(Subject).filter(Subject.id == semester.subject_id).first() if semester else None
    grade = db.query(Grade).filter(Grade.id == subject.grade_id).first() if subject else None
    
    exam_points = db.query(ExamPoint).filter(ExamPoint.unit_id == unit_id).all()
    
    doc = generate_unit_word_doc(
        unit=unit,
        semester=semester,
        subject=subject,
        grade=grade,
        unit_knowledge_json=unit.unit_knowledge_json,
        exam_points=exam_points
    )
    
    title_text = f"{grade.name if grade else ''} {subject.name if subject else ''} {semester.name if semester else ''} - {unit.name}"
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"{title_text}.docx"
    encoded_filename = quote(filename)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"{encoded_filename}\""}
    )

@router.post("/download-paper/{unit_id}")
def download_paper(
    unit_id: int,
    data: dict,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    question_count = data.get('question_count', 10)
    POINTS_COST = question_count
    
    if user.points < POINTS_COST:
        raise HTTPException(status_code=400, detail="积分不足")
    
    unit = db.query(Unit).filter(Unit.id == unit_id).first()
    if not unit:
        raise HTTPException(status_code=404, detail="单元不存在")
    
    user.points -= POINTS_COST
    
    semester = db.query(Semester).filter(Semester.id == unit.semester_id).first()
    subject = db.query(Subject).filter(Subject.id == semester.subject_id).first() if semester else None
    grade = db.query(Grade).filter(Grade.id == subject.grade_id).first() if subject else None
    
    title = f"{grade.name if grade else ''} {subject.name if subject else ''} {semester.name if semester else ''} - {unit.name} - 单元测试卷（{question_count}题）"
    
    order = Order(
        user_id=user.id,
        title=title,
        order_type="practice",
        points=POINTS_COST
    )
    db.add(order)
    db.commit()
    
    return {"message": "下载成功", "points": user.points}

@router.post("/download-unit/{unit_id}")
def download_unit(
    unit_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    POINTS_COST = 10
    
    if user.points < POINTS_COST:
        raise HTTPException(status_code=400, detail="积分不足")
    
    unit = db.query(Unit).filter(Unit.id == unit_id).first()
    if not unit:
        raise HTTPException(status_code=404, detail="单元不存在")
    
    user.points -= POINTS_COST
    
    semester = db.query(Semester).filter(Semester.id == unit.semester_id).first()
    subject = db.query(Subject).filter(Subject.id == semester.subject_id).first() if semester else None
    grade = db.query(Grade).filter(Grade.id == subject.grade_id).first() if subject else None
    
    title = f"{grade.name if grade else ''} {subject.name if subject else ''} {semester.name if semester else ''} - {unit.name} - 知识考点"
    
    order = Order(
        user_id=user.id,
        title=title,
        order_type="knowledge",
        points=POINTS_COST
    )
    db.add(order)
    db.commit()
    
    return {"message": "下载成功", "points": user.points}

@router.post("/set-child-grade")
def set_child_grade(
    data: dict,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    grade_id = data.get('grade_id')
    if not grade_id:
        raise HTTPException(status_code=400, detail="年级ID不能为空")
    
    grade = db.query(Grade).filter(Grade.id == grade_id).first()
    if not grade:
        raise HTTPException(status_code=404, detail="年级不存在")
    
    from datetime import datetime
    user.child_grade_id = grade_id
    user.child_grade_set_at = datetime.now()
    db.commit()
    
    return {"message": "设置成功", "grade_name": grade.name}

@router.get("/unit-question-stats/{unit_id}")
def get_unit_question_stats(
    unit_id: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    unit = db.query(Unit).filter(Unit.id == unit_id).first()
    if not unit:
        raise HTTPException(status_code=404, detail="单元不存在")
    
    total = db.query(Question).filter(Question.unit_id == unit_id).count()
    
    difficulties = db.query(Difficulty).all()
    difficulty_stats = []
    for d in difficulties:
        count = db.query(Question).filter(
            Question.unit_id == unit_id,
            Question.difficulty_id == d.id
        ).count()
        difficulty_stats.append({
            "id": d.id,
            "name": d.name,
            "count": count
        })
    
    question_types = db.query(QuestionType).all()
    type_stats = []
    for t in question_types:
        count = db.query(Question).filter(
            Question.unit_id == unit_id,
            Question.question_type_id == t.id
        ).count()
        type_stats.append({
            "id": t.id,
            "name": t.name,
            "count": count
        })
    
    return {
        "unit_id": unit_id,
        "unit_name": unit.name,
        "total": total,
        "difficulty_stats": difficulty_stats,
        "type_stats": type_stats
    }

@router.get("/generate-paper/{unit_id}")
def generate_paper(
    unit_id: int,
    question_count: int = 10,
    difficulty_id: int = None,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    unit = db.query(Unit).filter(Unit.id == unit_id).first()
    if not unit:
        raise HTTPException(status_code=404, detail="单元不存在")
    
    query = db.query(Question).filter(Question.unit_id == unit_id)
    if difficulty_id:
        query = query.filter(Question.difficulty_id == difficulty_id)
    
    questions = query.all()
    
    if not questions:
        raise HTTPException(status_code=400, detail="该单元暂无题目")
    
    if question_count > len(questions):
        question_count = len(questions)
    
    selected = random.sample(questions, question_count)
    
    result = []
    for idx, q in enumerate(selected, 1):
        options_list = []
        if q.options:
            try:
                import json
                options_list = json.loads(q.options) if isinstance(q.options, str) else q.options
            except:
                options_list = []
        
        result.append({
            "number": idx,
            "id": q.id,
            "content": q.content,
            "options": options_list,
            "answer": q.answer,
            "analysis": q.analysis,
            "question_type": q.question_type_obj.name if q.question_type_obj else "",
            "difficulty": q.difficulty_obj.name if q.difficulty_obj else "",
            "question_json": q.question_json,
            "stem": q.stem
        })
    
    return {
        "unit_id": unit_id,
        "unit_name": unit.name,
        "total": len(result),
        "questions": result
    }

def generate_paper_word_doc(unit, semester, subject, grade, questions):
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    title_text = f"{grade.name if grade else ''} {subject.name if subject else ''} {semester.name if semester else ''} - {unit.name} 单元测试卷"
    
    title = doc.add_paragraph()
    title_run = title.add_run(title_text)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = '宋体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    questions_by_type = {}
    for q in questions:
        q_type = q.get('question_type', '')
        if q_type not in questions_by_type:
            questions_by_type[q_type] = []
        questions_by_type[q_type].append(q)
    
    roman_numerals = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
    
    type_index = 0
    for q_type, q_list in questions_by_type.items():
        if type_index >= len(roman_numerals):
            type_num = str(type_index + 1)
        else:
            type_num = roman_numerals[type_index]
        
        section_title = doc.add_paragraph()
        section_run = section_title.add_run(f"{type_num}、{q_type}（共{len(q_list)}题）")
        section_run.font.size = Pt(14)
        section_run.font.bold = True
        section_run.font.name = '宋体'
        section_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        section_run.font.color.rgb = RGBColor(0, 102, 204)
        
        doc.add_paragraph()
        
        for idx, q in enumerate(q_list, 1):
            q_content = q.get('content', '')
            q_options = q.get('options', [])
            
            p = doc.add_paragraph()
            num_run = p.add_run(f"{idx}. ")
            num_run.font.bold = True
            num_run.font.name = '宋体'
            num_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            content_run = p.add_run(q_content)
            content_run.font.name = '宋体'
            content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            if q_options and isinstance(q_options, list):
                for opt_idx, opt in enumerate(q_options):
                    opt_letter = chr(65 + opt_idx)
                    opt_p = doc.add_paragraph()
                    opt_p.paragraph_format.left_indent = Pt(36)
                    opt_run = opt_p.add_run(f"{opt_letter}. {opt}")
                    opt_run.font.name = '宋体'
                    opt_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            doc.add_paragraph()
        
        type_index += 1
    
    doc.add_page_break()
    
    answer_title = doc.add_paragraph()
    answer_run = answer_title.add_run("参考答案及解析")
    answer_run.font.size = Pt(16)
    answer_run.font.bold = True
    answer_run.font.name = '宋体'
    answer_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    type_index = 0
    for q_type, q_list in questions_by_type.items():
        if type_index >= len(roman_numerals):
            type_num = str(type_index + 1)
        else:
            type_num = roman_numerals[type_index]
        
        section_title = doc.add_paragraph()
        section_run = section_title.add_run(f"{type_num}、{q_type}")
        section_run.font.size = Pt(14)
        section_run.font.bold = True
        section_run.font.name = '宋体'
        section_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        section_run.font.color.rgb = RGBColor(0, 102, 204)
        
        doc.add_paragraph()
        
        for idx, q in enumerate(q_list, 1):
            q_answer = q.get('answer', '')
            q_analysis = q.get('analysis', '')
            
            ans_p = doc.add_paragraph()
            ans_num_run = ans_p.add_run(f"{idx}. ")
            ans_num_run.font.bold = True
            ans_num_run.font.name = '宋体'
            ans_num_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            ans_label_run = ans_p.add_run("答案：")
            ans_label_run.font.bold = True
            ans_label_run.font.color.rgb = RGBColor(220, 20, 60)
            ans_label_run.font.name = '宋体'
            ans_label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            ans_content_run = ans_p.add_run(q_answer)
            ans_content_run.font.name = '宋体'
            ans_content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            if q_analysis:
                analysis_p = doc.add_paragraph()
                analysis_p.paragraph_format.left_indent = Pt(24)
                analysis_label_run = analysis_p.add_run("解析：")
                analysis_label_run.font.bold = True
                analysis_label_run.font.color.rgb = RGBColor(0, 139, 69)
                analysis_label_run.font.name = '宋体'
                analysis_label_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                
                analysis_content_run = analysis_p.add_run(q_analysis)
                analysis_content_run.font.name = '宋体'
                analysis_content_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            doc.add_paragraph()
        
        type_index += 1
    
    return doc

@router.get("/paper-word/{unit_id}")
def get_paper_word(
    unit_id: int,
    question_count: int = 10,
    request: Request,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    unit = db.query(Unit).filter(Unit.id == unit_id).first()
    if not unit:
        raise HTTPException(status_code=404, detail="单元不存在")
    
    semester = db.query(Semester).filter(Semester.id == unit.semester_id).first()
    subject = db.query(Subject).filter(Subject.id == semester.subject_id).first() if semester else None
    grade = db.query(Grade).filter(Grade.id == subject.grade_id).first() if subject else None
    
    difficulty_counts = {}
    question_type_counts = {}
    
    query_params = request.query_params
    for key in query_params.keys():
        value = query_params[key]
        if key.startswith('difficulty_'):
            try:
                diff_id = int(key.split('_')[1])
                difficulty_counts[diff_id] = int(value)
            except:
                pass
        elif key.startswith('type_'):
            try:
                type_id = int(key.split('_')[1])
                question_type_counts[type_id] = int(value)
            except:
                pass
    
    all_selected = []
    
    if question_type_counts:
        for type_id, count in question_type_counts.items():
            if count <= 0:
                continue
            
            query = db.query(Question).filter(
                Question.unit_id == unit_id,
                Question.question_type_id == type_id
            )
            
            if difficulty_counts:
                query = query.filter(Question.difficulty_id.in_(difficulty_counts.keys()))
            
            questions = query.all()
            
            if not questions:
                continue
            
            selected_count = min(count, len(questions))
            selected = random.sample(questions, selected_count)
            
            for q in selected:
                options_list = []
                if q.options:
                    try:
                        import json
                        options_list = json.loads(q.options) if isinstance(q.options, str) else q.options
                    except:
                        options_list = []
                
                all_selected.append({
                    "content": q.content,
                    "options": options_list,
                    "answer": q.answer,
                    "analysis": q.analysis,
                    "question_type": q.question_type_obj.name if q.question_type_obj else "",
                    "difficulty": q.difficulty_obj.name if q.difficulty_obj else "",
                    "question_type_id": q.question_type_id,
                })
    else:
        query = db.query(Question).filter(Question.unit_id == unit_id)
        
        if difficulty_counts:
            query = query.filter(Question.difficulty_id.in_(difficulty_counts.keys()))
        
        questions = query.all()
        
        if not questions:
            raise HTTPException(status_code=400, detail="该单元暂无符合条件的题目")
        
        selected_count = min(question_count, len(questions))
        selected = random.sample(questions, selected_count)
        
        for q in selected:
            options_list = []
            if q.options:
                try:
                    import json
                    options_list = json.loads(q.options) if isinstance(q.options, str) else q.options
                except:
                    options_list = []
            
            all_selected.append({
                "content": q.content,
                "options": options_list,
                "answer": q.answer,
                "analysis": q.analysis,
                "question_type": q.question_type_obj.name if q.question_type_obj else "",
                "difficulty": q.difficulty_obj.name if q.difficulty_obj else "",
                "question_type_id": q.question_type_id,
            })
    
    doc = generate_paper_word_doc(unit, semester, subject, grade, all_selected)
    
    title_text = f"{grade.name if grade else ''} {subject.name if subject else ''} {semester.name if semester else ''} - {unit.name} 单元测试卷"
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"{title_text}.docx"
    encoded_filename = quote(filename)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"{encoded_filename}\""}
    )