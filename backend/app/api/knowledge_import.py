from fastapi import APIRouter, Depends, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from app.database import get_db
from app.models.models import (
    User, Version, Grade, Subject, Semester, Unit,
    KnowledgePoint, ExamPoint, Question, QuestionType, Difficulty
)
from app.utils.auth import get_current_admin
from openpyxl import load_workbook
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from typing import List
from urllib.parse import quote

router = APIRouter(prefix="/admin", tags=["知识考点管理"])

def set_run_shading(run, color='D9D9D9'):
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    rPr.append(shd)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def format_content(doc, content):
    if not content:
        return
    
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        
        if line.startswith('【') and '】' in line:
            end_idx = line.index('】')
            title_text = line[:end_idx + 1]
            rest_text = line[end_idx + 1:].strip()
            
            p = doc.add_paragraph()
            run = p.add_run(title_text)
            run.font.bold = True
            run.font.size = Pt(12)
            
            if rest_text:
                p2 = doc.add_paragraph()
                p2.paragraph_format.first_line_indent = Pt(24)
                p2.add_run(rest_text)
            
            i += 1
        elif '：' in line or ':' in line:
            colon_char = '：' if '：' in line else ':'
            colon_idx = line.index(colon_char)
            key_text = line[:colon_idx].strip()
            value_text = line[colon_idx + 1:].strip()
            
            p = doc.add_paragraph()
            key_run = p.add_run(key_text + colon_char)
            key_run.font.bold = True
            set_run_shading(key_run)
            if value_text:
                p.add_run(value_text)
            i += 1
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)
            p.add_run(line)
            i += 1

@router.get("/knowledge-exam-points")
def get_knowledge_exam_points(unit_id: int = None, admin: User = Depends(get_current_admin), db: Session = Depends(get_db)):
    query = db.query(KnowledgePoint)
    if unit_id:
        query = query.filter(KnowledgePoint.unit_id == unit_id)
    
    knowledge_points = query.all()
    result = []
    for kp in knowledge_points:
        exam_points = db.query(ExamPoint).filter(ExamPoint.unit_id == kp.unit_id).all()
        result.append({
            "id": kp.id,
            "unit_id": kp.unit_id,
            "title": kp.title,
            "content": kp.content,
            "exam_points": [{
                "id": ep.id,
                "title": ep.title,
                "content": ep.content,
                "exam_types": ep.exam_types,
                "exam_frequency": ep.exam_frequency.value if ep.exam_frequency else None
            } for ep in exam_points]
        })
    return result

@router.post("/knowledge-exam-points")
def save_knowledge_exam_point(data: dict, admin: User = Depends(get_current_admin), db: Session = Depends(get_db)):
    if data.get('id'):
        kp = db.query(KnowledgePoint).filter(KnowledgePoint.id == data['id']).first()
        if kp:
            kp.title = data.get('title', kp.title)
            kp.content = data.get('content', kp.content)
            db.commit()
            
            if 'exam_points' in data:
                for ep_data in data['exam_points']:
                    if ep_data.get('id'):
                        ep = db.query(ExamPoint).filter(ExamPoint.id == ep_data['id']).first()
                        if ep:
                            ep.title = ep_data.get('title', ep.title)
                            ep.content = ep_data.get('content', ep.content)
                            ep.exam_types = ep_data.get('exam_types', ep.exam_types)
                            ep.exam_frequency = ep_data.get('exam_frequency', ep.exam_frequency)
                    else:
                        ep = ExamPoint(
                            knowledge_point_id=kp.id,
                            title=ep_data.get('title'),
                            content=ep_data.get('content'),
                            exam_types=ep_data.get('exam_types'),
                            exam_frequency=ep_data.get('exam_frequency', '常考')
                        )
                        db.add(ep)
                db.commit()
            return {"message": "更新成功"}
    
    kp = KnowledgePoint(
        unit_id=data['unit_id'],
        title=data['title'],
        content=data.get('content')
    )
    db.add(kp)
    db.commit()
    db.refresh(kp)
    
    if 'exam_points' in data:
        for ep_data in data['exam_points']:
            ep = ExamPoint(
                knowledge_point_id=kp.id,
                title=ep_data.get('title'),
                content=ep_data.get('content'),
                exam_types=ep_data.get('exam_types'),
                exam_frequency=ep_data.get('exam_frequency', '常考')
            )
            db.add(ep)
        db.commit()
    
    return {"id": kp.id, "message": "保存成功"}

@router.post("/import-knowledge-exam-points")
async def import_knowledge_exam_points(
    file: UploadFile = File(...),
    version_id: int = None,
    admin: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="只支持Excel文件")
    
    contents = await file.read()
    try:
        wb = load_workbook(BytesIO(contents))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Excel文件解析失败: {str(e)}")
    
    imported_count = 0
    skipped_count = 0
    duplicate_count = 0
    unit_knowledge_map = {}
    import_log = []
    
    version = db.query(Version).filter(Version.name == "人教版").first()
    if not version:
        version = Version(name="人教版")
        db.add(version)
        db.commit()
        db.refresh(version)
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        
        subject_name = sheet_name.replace('小学', '').replace('初中', '').strip()
        import_log.append(f"📄 处理Sheet: {sheet_name} -> 科目: {subject_name}")
        
        last_grade = None
        last_semester = None
        last_unit_number = None
        last_unit_name = None
        last_knowledge_content = None
        
        for row_idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
            if not row:
                continue
            
            grade_name = str(row[0]).strip() if row[0] else None
            semester_name = str(row[1]).strip() if row[1] else None
            unit_number_text = str(row[2]).strip() if row[2] else None
            unit_name = str(row[3]).strip() if row[3] else None
            knowledge_content = str(row[4]).strip() if row[4] else None
            exam_content = str(row[5]).strip() if row[5] else None
            exam_types = str(row[6]).strip() if row[6] else None
            exam_frequency = str(row[7]).strip() if row[7] else '常考'
            
            if grade_name:
                last_grade = grade_name
            if semester_name:
                last_semester = semester_name
            if unit_number_text:
                chinese_nums = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}
                for cn, num in chinese_nums.items():
                    if cn in unit_number_text:
                        last_unit_number = num
                        break
                if not last_unit_number:
                    import re
                    match = re.search(r'\d+', unit_number_text)
                    if match:
                        last_unit_number = int(match.group())
            if unit_name:
                last_unit_name = unit_name
            if knowledge_content:
                last_knowledge_content = knowledge_content
            
            if not all([last_grade, last_semester, last_unit_number, last_unit_name]):
                skipped_count += 1
                import_log.append(f"  ✗ 行{row_idx}: 缺少必要字段")
                continue
            
            if not exam_content:
                skipped_count += 1
                import_log.append(f"  ✗ 行{row_idx}: {last_unit_name} 无考点内容")
                continue
            
            grade = db.query(Grade).filter(
                Grade.version_id == version.id,
                Grade.name == last_grade
            ).first()
            if not grade:
                grade = Grade(version_id=version.id, name=last_grade)
                db.add(grade)
                db.commit()
                db.refresh(grade)
            
            subject = db.query(Subject).filter(
                Subject.grade_id == grade.id,
                Subject.name == subject_name
            ).first()
            if not subject:
                subject = Subject(grade_id=grade.id, name=subject_name)
                db.add(subject)
                db.commit()
                db.refresh(subject)
            
            semester = db.query(Semester).filter(
                Semester.subject_id == subject.id,
                Semester.name == last_semester
            ).first()
            if not semester:
                semester = Semester(subject_id=subject.id, name=last_semester)
                db.add(semester)
                db.commit()
                db.refresh(semester)
            
            unit = db.query(Unit).filter(
                Unit.semester_id == semester.id,
                Unit.unit_number == last_unit_number
            ).first()
            if not unit:
                unit = Unit(
                    semester_id=semester.id,
                    unit_number=last_unit_number,
                    name=f"{unit_number_text} {last_unit_name}"
                )
                db.add(unit)
                db.commit()
                db.refresh(unit)
            
            unit_key = f"{last_grade}-{subject_name}-{last_semester}-{last_unit_number}"
            
            if unit_key not in unit_knowledge_map:
                kp = db.query(KnowledgePoint).filter(
                    KnowledgePoint.unit_id == unit.id
                ).first()
                
                if not kp:
                    kp_content = last_knowledge_content or ""
                    kp = KnowledgePoint(
                        unit_id=unit.id,
                        title=f"{last_unit_name} - 知识点",
                        content=kp_content
                    )
                    db.add(kp)
                    db.commit()
                    db.refresh(kp)
                    import_log.append(f"  ✓ 创建知识点: {last_unit_name}")
                
                unit_knowledge_map[unit_key] = kp
            
            kp = unit_knowledge_map[unit_key]
            
            if last_knowledge_content and kp.content != last_knowledge_content:
                kp.content = last_knowledge_content
                db.commit()
            
            if exam_content:
                existing_ep = db.query(ExamPoint).filter(
                    ExamPoint.unit_id == unit.id,
                    ExamPoint.content == exam_content
                ).first()
                
                if not existing_ep:
                    if '：' in exam_content:
                        colon_idx = exam_content.index('：')
                        exam_title = exam_content[:colon_idx].strip()
                        exam_content_clean = exam_content[colon_idx + 1:].strip()
                    elif ':' in exam_content:
                        colon_idx = exam_content.index(':')
                        exam_title = exam_content[:colon_idx].strip()
                        exam_content_clean = exam_content[colon_idx + 1:].strip()
                    else:
                        lines = exam_content.split('\n')
                        exam_title = lines[0][:50] if lines else exam_content[:50]
                        exam_content_clean = '\n'.join(lines[1:]) if len(lines) > 1 else exam_content
                    
                    freq = '必考' if exam_frequency in ['必考', '必考重点'] else ('常考' if exam_frequency == '常考' else '少考')
                    ep = ExamPoint(
                        unit_id=unit.id,
                        title=exam_title,
                        content=exam_content_clean,
                        exam_types=exam_types,
                        exam_frequency=freq
                    )
                    db.add(ep)
                    db.commit()
                    imported_count += 1
                    import_log.append(f"  ✓ 行{row_idx}: {exam_title[:30]} ({freq})")
                else:
                    duplicate_count += 1
                    import_log.append(f"  - 行{row_idx}: 重复考点")
    
    return {
        "message": f"导入完成",
        "imported": imported_count,
        "skipped": skipped_count,
        "duplicate": duplicate_count,
        "log": import_log[:50]
    }

@router.post("/import-knowledge")
async def import_knowledge(
    file: UploadFile = File(...),
    version_id: int = None,
    grade_id: int = None,
    subject_id: int = None,
    semester_id: int = None,
    admin: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="只支持Excel文件")
    
    contents = await file.read()
    try:
        wb = load_workbook(BytesIO(contents))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Excel文件解析失败: {str(e)}")
    
    imported_count = 0
    import_log = []
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        import_log.append(f"📄 处理Sheet: {sheet_name}")
        
        for row_idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
            if not row or not row[0]:
                continue
            
            subject_name = str(row[0]).strip() if row[0] else None
            grade_name = str(row[1]).strip() if row[1] else None
            semester_name = str(row[2]).strip() if row[2] else None
            unit_number = int(row[3]) if row[3] else None
            unit_number_text = str(row[4]).strip() if row[4] else None
            unit_name = str(row[5]).strip() if row[5] else None
            knowledge_content = str(row[6]).strip() if len(row) > 6 and row[6] else None
            
            if not all([subject_name, grade_name, semester_name, unit_number, unit_name, knowledge_content]):
                import_log.append(f"  ✗ 行{row_idx}: 缺少必要字段")
                continue
            
            grade_obj = db.query(Grade).filter(Grade.name == grade_name).first()
            if not grade_obj:
                import_log.append(f"  ✗ 行{row_idx}: 年级'{grade_name}'不存在")
                continue
            if grade_id and grade_obj.id != grade_id:
                continue
            
            subject_obj = db.query(Subject).filter(
                Subject.grade_id == grade_obj.id,
                Subject.name == subject_name
            ).first()
            if not subject_obj:
                import_log.append(f"  ✗ 行{row_idx}: 科目'{subject_name}'不存在")
                continue
            if subject_id and subject_obj.id != subject_id:
                continue
            
            semester_obj = db.query(Semester).filter(
                Semester.subject_id == subject_obj.id,
                Semester.name == semester_name
            ).first()
            if not semester_obj:
                import_log.append(f"  ✗ 行{row_idx}: 学期'{semester_name}'不存在")
                continue
            if semester_id and semester_obj.id != semester_id:
                continue
            
            unit = db.query(Unit).filter(
                Unit.semester_id == semester_obj.id,
                Unit.unit_number == unit_number
            ).first()
            if not unit:
                import_log.append(f"  ✗ 行{row_idx}: 单元序号{unit_number}不存在")
                continue
            
            kp = db.query(KnowledgePoint).filter(KnowledgePoint.unit_id == unit.id).first()
            if kp:
                db.delete(kp)
                db.commit()
            
            kp = KnowledgePoint(
                unit_id=unit.id,
                title=f"{unit_name} - 知识点",
                content=knowledge_content
            )
            db.add(kp)
            db.commit()
            imported_count += 1
            import_log.append(f"  ✓ {grade_name} {subject_name} {semester_name} - {unit_name}")
    
    return {"message": f"导入完成，共{imported_count}个知识点", "log": import_log[:50]}

@router.post("/clear-knowledge")
def clear_knowledge(
    data: dict,
    admin: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    version_id = data.get('version_id')
    grade_id = data.get('grade_id')
    subject_id = data.get('subject_id')
    semester_id = data.get('semester_id')
    
    query = db.query(KnowledgePoint)
    
    if semester_id:
        unit_ids = [u.id for u in db.query(Unit).filter(Unit.semester_id == semester_id).all()]
        query = query.filter(KnowledgePoint.unit_id.in_(unit_ids))
    elif subject_id:
        semester_ids = [s.id for s in db.query(Semester).filter(Semester.subject_id == subject_id).all()]
        unit_ids = [u.id for u in db.query(Unit).filter(Unit.semester_id.in_(semester_ids)).all()]
        query = query.filter(KnowledgePoint.unit_id.in_(unit_ids))
    elif grade_id:
        subject_ids = [s.id for s in db.query(Subject).filter(Subject.grade_id == grade_id).all()]
        semester_ids = [s.id for s in db.query(Semester).filter(Semester.subject_id.in_(subject_ids)).all()]
        unit_ids = [u.id for u in db.query(Unit).filter(Unit.semester_id.in_(semester_ids)).all()]
        query = query.filter(KnowledgePoint.unit_id.in_(unit_ids))
    elif version_id:
        grade_ids = [g.id for g in db.query(Grade).filter(Grade.version_id == version_id).all()]
        subject_ids = [s.id for s in db.query(Subject).filter(Subject.grade_id.in_(grade_ids)).all()]
        semester_ids = [s.id for s in db.query(Semester).filter(Semester.subject_id.in_(subject_ids)).all()]
        unit_ids = [u.id for u in db.query(Unit).filter(Unit.semester_id.in_(semester_ids)).all()]
        query = query.filter(KnowledgePoint.unit_id.in_(unit_ids))
    
    count = query.count()
    query.delete()
    db.commit()
    
    return {"message": f"已清除{count}个知识点"}

@router.post("/import-exam-points")
async def import_exam_points(
    file: UploadFile = File(...),
    version_id: int = None,
    grade_id: int = None,
    subject_id: int = None,
    semester_id: int = None,
    admin: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="只支持Excel文件")
    
    contents = await file.read()
    try:
        wb = load_workbook(BytesIO(contents))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Excel文件解析失败: {str(e)}")
    
    imported_count = 0
    import_log = []
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        import_log.append(f"📄 处理Sheet: {sheet_name}")
        
        for row_idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
            if not row or not row[0]:
                continue
            
            subject_name = str(row[0]).strip() if row[0] else None
            grade_name = str(row[1]).strip() if row[1] else None
            semester_name = str(row[2]).strip() if row[2] else None
            unit_number = int(row[3]) if row[3] else None
            unit_number_text = str(row[4]).strip() if row[4] else None
            unit_name = str(row[5]).strip() if row[5] else None
            exam_content = str(row[6]).strip() if len(row) > 6 and row[6] else None
            exam_types = str(row[7]).strip() if len(row) > 7 and row[7] else None
            exam_frequency = '常考'
            
            if not all([subject_name, grade_name, semester_name, unit_number, exam_content]):
                import_log.append(f"  ✗ 行{row_idx}: 缺少必要字段")
                continue
            
            grade_obj = db.query(Grade).filter(Grade.name == grade_name).first()
            if not grade_obj:
                import_log.append(f"  ✗ 行{row_idx}: 年级'{grade_name}'不存在")
                continue
            if grade_id and grade_obj.id != grade_id:
                continue
            
            subject_obj = db.query(Subject).filter(
                Subject.grade_id == grade_obj.id,
                Subject.name == subject_name
            ).first()
            if not subject_obj:
                import_log.append(f"  ✗ 行{row_idx}: 科目'{subject_name}'不存在")
                continue
            if subject_id and subject_obj.id != subject_id:
                continue
            
            semester_obj = db.query(Semester).filter(
                Semester.subject_id == subject_obj.id,
                Semester.name == semester_name
            ).first()
            if not semester_obj:
                import_log.append(f"  ✗ 行{row_idx}: 学期'{semester_name}'不存在")
                continue
            if semester_id and semester_obj.id != semester_id:
                continue
            
            unit = db.query(Unit).filter(
                Unit.semester_id == semester_obj.id,
                Unit.unit_number == unit_number
            ).first()
            if not unit:
                import_log.append(f"  ✗ 行{row_idx}: 单元序号{unit_number}不存在")
                continue
            
            existing_ep = db.query(ExamPoint).filter(
                ExamPoint.unit_id == unit.id,
                ExamPoint.content == exam_content
            ).first()
            
            if existing_ep:
                db.delete(existing_ep)
                db.commit()
            
            if '：' in exam_content:
                colon_idx = exam_content.index('：')
                exam_title = exam_content[:colon_idx].strip()
                exam_content_clean = exam_content[colon_idx + 1:].strip()
            elif ':' in exam_content:
                colon_idx = exam_content.index(':')
                exam_title = exam_content[:colon_idx].strip()
                exam_content_clean = exam_content[colon_idx + 1:].strip()
            else:
                lines = exam_content.split('\n')
                exam_title = lines[0][:50] if lines else exam_content[:50]
                exam_content_clean = '\n'.join(lines[1:]) if len(lines) > 1 else exam_content
            
            freq = '必考' if exam_frequency in ['必考', '必考重点'] else ('常考' if exam_frequency == '常考' else '少考')
            
            ep = ExamPoint(
                unit_id=unit.id,
                title=exam_title,
                content=exam_content_clean,
                exam_types=exam_types,
                exam_frequency=freq
            )
            db.add(ep)
            db.commit()
            imported_count += 1
            import_log.append(f"  ✓ {grade_name} {subject_name} {semester_name} - {exam_title[:30]}")
    
    return {"message": f"导入完成，共{imported_count}个考点", "log": import_log[:50]}

@router.post("/clear-exam-points")
def clear_exam_points(
    data: dict,
    admin: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    version_id = data.get('version_id')
    grade_id = data.get('grade_id')
    subject_id = data.get('subject_id')
    semester_id = data.get('semester_id')
    
    query = db.query(ExamPoint)
    
    if semester_id:
        unit_ids = [u.id for u in db.query(Unit).filter(Unit.semester_id == semester_id).all()]
        query = query.filter(ExamPoint.unit_id.in_(unit_ids))
    elif subject_id:
        semester_ids = [s.id for s in db.query(Semester).filter(Semester.subject_id == subject_id).all()]
        unit_ids = [u.id for u in db.query(Unit).filter(Unit.semester_id.in_(semester_ids)).all()]
        query = query.filter(ExamPoint.unit_id.in_(unit_ids))
    elif grade_id:
        subject_ids = [s.id for s in db.query(Subject).filter(Subject.grade_id == grade_id).all()]
        semester_ids = [s.id for s in db.query(Semester).filter(Semester.subject_id.in_(subject_ids)).all()]
        unit_ids = [u.id for u in db.query(Unit).filter(Unit.semester_id.in_(semester_ids)).all()]
        query = query.filter(ExamPoint.unit_id.in_(unit_ids))
    elif version_id:
        grade_ids = [g.id for g in db.query(Grade).filter(Grade.version_id == version_id).all()]
        subject_ids = [s.id for s in db.query(Subject).filter(Subject.grade_id.in_(grade_ids)).all()]
        semester_ids = [s.id for s in db.query(Semester).filter(Semester.subject_id.in_(subject_ids)).all()]
        unit_ids = [u.id for u in db.query(Unit).filter(Unit.semester_id.in_(semester_ids)).all()]
        query = query.filter(ExamPoint.unit_id.in_(unit_ids))
    
    count = query.count()
    query.delete()
    db.commit()
    
    return {"message": f"已清除{count}个考点"}

@router.post("/clean-duplicate-exam-points")
def clean_duplicate_exam_points(admin: User = Depends(get_current_admin), db: Session = Depends(get_db)):
    all_units = db.query(Unit).all()
    deleted_count = 0
    
    for unit in all_units:
        exam_points = db.query(ExamPoint).filter(ExamPoint.unit_id == unit.id).all()
        
        seen_contents = set()
        for ep in exam_points:
            if ep.content in seen_contents:
                db.delete(ep)
                deleted_count += 1
            else:
                seen_contents.add(ep.content)
    
    db.commit()
    return {"message": f"清理完成，删除了{deleted_count}条重复考点"}

@router.delete("/unit-knowledge/{unit_id}")
def delete_unit_knowledge(
    unit_id: int,
    admin: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    db.query(ExamPoint).filter(ExamPoint.unit_id == unit_id).delete()
    kps = db.query(KnowledgePoint).filter(KnowledgePoint.unit_id == unit_id).all()
    deleted_count = 0
    for kp in kps:
        db.delete(kp)
        deleted_count += 1
    db.commit()
    return {"message": f"删除成功，共删除{deleted_count}个知识点及其考点"}

@router.post("/clean-exam-content")
def clean_exam_content(admin: User = Depends(get_current_admin), db: Session = Depends(get_db)):
    all_eps = db.query(ExamPoint).all()
    cleaned_count = 0
    
    for ep in all_eps:
        if ep.content and ep.title:
            if '：' in ep.content:
                colon_idx = ep.content.index('：')
                new_title = ep.content[:colon_idx].strip()
                new_content = ep.content[colon_idx + 1:].strip()
                if new_title != ep.title or new_content != ep.content:
                    ep.title = new_title
                    ep.content = new_content
                    cleaned_count += 1
            elif ':' in ep.content:
                colon_idx = ep.content.index(':')
                new_title = ep.content[:colon_idx].strip()
                new_content = ep.content[colon_idx + 1:].strip()
                if new_title != ep.title or new_content != ep.content:
                    ep.title = new_title
                    ep.content = new_content
                    cleaned_count += 1
    
    db.commit()
    return {"message": f"清理完成，修复了{cleaned_count}条考点内容"}


@router.post("/import-units")
async def import_units(
    file: UploadFile = File(...),
    admin: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="只支持Excel文件")
    
    contents = await file.read()
    try:
        wb = load_workbook(BytesIO(contents))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Excel文件解析失败: {str(e)}")
    
    imported_count = 0
    skipped_count = 0
    import_log = []
    
    version = db.query(Version).filter(Version.name == "人教版").first()
    if not version:
        version = Version(name="人教版")
        db.add(version)
        db.commit()
        db.refresh(version)
    
    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        import_log.append(f"📄 处理Sheet: {sheet_name}")
        
        for row_idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
            if not row or not row[0]:
                continue
            
            subject_name = str(row[0]).strip() if row[0] else None
            grade_name = str(row[1]).strip() if row[1] else None
            semester_name = str(row[2]).strip() if row[2] else None
            unit_number = int(row[3]) if row[3] else None
            unit_number_text = str(row[4]).strip() if row[4] else None
            unit_name = str(row[5]).strip() if row[5] else None
            
            if not all([subject_name, grade_name, semester_name, unit_number, unit_number_text, unit_name]):
                skipped_count += 1
                import_log.append(f"  ✗ 行{row_idx}: 缺少必要字段")
                continue
            
            grade = db.query(Grade).filter(
                Grade.version_id == version.id,
                Grade.name == grade_name
            ).first()
            if not grade:
                grade = Grade(version_id=version.id, name=grade_name)
                db.add(grade)
                db.commit()
                db.refresh(grade)
            
            subject = db.query(Subject).filter(
                Subject.grade_id == grade.id,
                Subject.name == subject_name
            ).first()
            if not subject:
                subject = Subject(grade_id=grade.id, name=subject_name)
                db.add(subject)
                db.commit()
                db.refresh(subject)
            
            semester = db.query(Semester).filter(
                Semester.subject_id == subject.id,
                Semester.name == semester_name
            ).first()
            if not semester:
                semester = Semester(subject_id=subject.id, name=semester_name)
                db.add(semester)
                db.commit()
                db.refresh(semester)
            
            unit = db.query(Unit).filter(
                Unit.semester_id == semester.id,
                Unit.unit_number == unit_number
            ).first()
            
            if not unit:
                unit = Unit(
                    semester_id=semester.id,
                    unit_number=unit_number,
                    name=f"{unit_number_text} {unit_name}"
                )
                db.add(unit)
                db.commit()
                db.refresh(unit)
                imported_count += 1
                import_log.append(f"  ✓ 行{row_idx}: {subject_name} {grade_name} {semester_name} - {unit_number_text} {unit_name}")
            else:
                old_name = unit.name
                unit.name = f"{unit_number_text} {unit_name}"
                db.commit()
                import_log.append(f"  ↻ 行{row_idx}: 更新 {old_name} → {unit.name}")
    
    return {
        "message": f"导入完成",
        "imported": imported_count,
        "skipped": skipped_count,
        "log": import_log[:50]
    }

@router.post("/query-units")
def query_units(
    data: dict,
    admin: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    version_id = data.get('version_id')
    grade_id = data.get('grade_id')
    subject_id = data.get('subject_id')
    semester_id = data.get('semester_id')
    
    query = db.query(Unit)
    
    if semester_id:
        query = query.filter(Unit.semester_id == semester_id)
    elif subject_id:
        semester_ids = [s.id for s in db.query(Semester).filter(Semester.subject_id == subject_id).all()]
        query = query.filter(Unit.semester_id.in_(semester_ids))
    elif grade_id:
        subject_ids = [s.id for s in db.query(Subject).filter(Subject.grade_id == grade_id).all()]
        semester_ids = [s.id for s in db.query(Semester).filter(Semester.subject_id.in_(subject_ids)).all()]
        query = query.filter(Unit.semester_id.in_(semester_ids))
    elif version_id:
        grade_ids = [g.id for g in db.query(Grade).filter(Grade.version_id == version_id).all()]
        subject_ids = [s.id for s in db.query(Subject).filter(Subject.grade_id.in_(grade_ids)).all()]
        semester_ids = [s.id for s in db.query(Semester).filter(Semester.subject_id.in_(subject_ids)).all()]
        query = query.filter(Unit.semester_id.in_(semester_ids))
    
    units = query.all()
    result = []
    
    for unit in units:
        semester = db.query(Semester).filter(Semester.id == unit.semester_id).first()
        subject = db.query(Subject).filter(Subject.id == semester.subject_id).first() if semester else None
        grade = db.query(Grade).filter(Grade.id == subject.grade_id).first() if subject else None
        
        has_knowledge = db.query(KnowledgePoint).filter(KnowledgePoint.unit_id == unit.id).first() is not None
        exam_point_count = db.query(ExamPoint).filter(ExamPoint.unit_id == unit.id).count()
        
        result.append({
            "id": unit.id,
            "grade_name": grade.name if grade else "",
            "subject_name": subject.name if subject else "",
            "semester_name": semester.name if semester else "",
            "unit_number": unit.unit_number,
            "name": unit.name,
            "has_knowledge": has_knowledge,
            "exam_point_count": exam_point_count
        })
    
    return {"units": result}

@router.get("/unit-detail/{unit_id}")
def get_unit_detail(
    unit_id: int,
    admin: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    knowledge = db.query(KnowledgePoint).filter(KnowledgePoint.unit_id == unit_id).first()
    exam_points = db.query(ExamPoint).filter(ExamPoint.unit_id == unit_id).all()
    
    return {
        "knowledge": {
            "id": knowledge.id,
            "title": knowledge.title,
            "content": knowledge.content
        } if knowledge else None,
        "exam_points": [{
            "id": ep.id,
            "title": ep.title,
            "content": ep.content,
            "exam_types": ep.exam_types,
            "exam_frequency": ep.exam_frequency.value if ep.exam_frequency else None
        } for ep in exam_points]
    }

@router.get("/unit-word/{unit_id}")
def get_unit_word(
    unit_id: int,
    admin: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    unit = db.query(Unit).filter(Unit.id == unit_id).first()
    if not unit:
        raise HTTPException(status_code=404, detail="单元不存在")
    
    semester = db.query(Semester).filter(Semester.id == unit.semester_id).first()
    subject = db.query(Subject).filter(Subject.id == semester.subject_id).first() if semester else None
    grade = db.query(Grade).filter(Grade.id == subject.grade_id).first() if subject else None
    
    knowledge = db.query(KnowledgePoint).filter(KnowledgePoint.unit_id == unit_id).first()
    exam_points = db.query(ExamPoint).filter(ExamPoint.unit_id == unit_id).all()
    
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    title_text = f"{grade.name if grade else ''} {subject.name if subject else ''} {semester.name if semester else ''} - {unit.name}"
    
    title = doc.add_paragraph()
    title_run = title.add_run(title_text)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = '宋体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("知识点")
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.bold = True
    subtitle_run.font.color.rgb = RGBColor(0, 102, 204)
    subtitle_run.font.name = '宋体'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if knowledge and knowledge.content:
        format_content(doc, knowledge.content)
    else:
        p = doc.add_paragraph()
        run = p.add_run("暂无知识点内容")
        run.font.color.rgb = RGBColor(153, 153, 153)
    
    doc.add_page_break()
    
    subtitle2 = doc.add_paragraph()
    subtitle2_run = subtitle2.add_run("考点")
    subtitle2_run.font.size = Pt(16)
    subtitle2_run.font.bold = True
    subtitle2_run.font.color.rgb = RGBColor(0, 102, 204)
    subtitle2_run.font.name = '宋体'
    subtitle2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if exam_points:
        for idx, ep in enumerate(exam_points, 1):
            p = doc.add_paragraph()
            
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(2.5), WD_TAB_ALIGNMENT.CENTER)
            tab_stops.add_tab_stop(Inches(5.5), WD_TAB_ALIGNMENT.RIGHT)
            
            title_run = p.add_run(f"{idx}. {ep.title}")
            title_run.font.size = Pt(14)
            title_run.font.bold = True
            title_run.font.name = '宋体'
            title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            p.add_run('\t')
            
            if ep.exam_types:
                types_run = p.add_run(f"{{{ep.exam_types}}}")
                types_run.font.size = Pt(12)
                types_run.font.name = '宋体'
                types_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                types_run.font.color.rgb = RGBColor(255, 105, 180)
            
            p.add_run('\t')
            
            if ep.exam_frequency:
                freq_text = ep.exam_frequency.value
                freq_run = p.add_run(f"[{freq_text}]")
                freq_run.font.size = Pt(12)
                freq_run.font.name = '宋体'
                freq_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                if freq_text == '必考':
                    freq_run.font.color.rgb = RGBColor(255, 0, 0)
                elif freq_text == '常考':
                    freq_run.font.color.rgb = RGBColor(255, 153, 0)
                else:
                    freq_run.font.color.rgb = RGBColor(0, 153, 0)
            
            if ep.content:
                format_content(doc, ep.content)
            
            doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        run = p.add_run("暂无考点内容")
        run.font.color.rgb = RGBColor(153, 153, 153)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"{title_text}.docx"
    encoded_filename = quote(filename)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=\"{encoded_filename}\""}
    )

@router.post("/import-questions")
async def import_questions(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    admin = Depends(get_current_admin)
):
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="只支持Excel文件")
    
    wb = load_workbook(BytesIO(await file.read()))
    ws = wb.active
    
    question_types_cache = {}
    difficulties_cache = {}
    
    imported = 0
    skipped = 0
    errors = []
    
    for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        if not row[0]:
            continue
        
        try:
            subject_name = row[0]
            grade_name = row[1]
            semester_name = row[2]
            unit_number = row[3]
            unit_name = row[5]
            exam_point_title = row[6]
            question_type_name = row[9]
            difficulty_name = row[10]
            stem_text = row[11]
            answer_text = row[12]
            analysis_text = row[13]
            json_content = row[14]
            
            grade = db.query(Grade).filter(Grade.name == grade_name).first()
            if not grade:
                skipped += 1
                continue
            
            subject = db.query(Subject).filter(
                Subject.grade_id == grade.id,
                Subject.name == subject_name
            ).first()
            if not subject:
                skipped += 1
                continue
            
            semester = db.query(Semester).filter(
                Semester.subject_id == subject.id,
                Semester.name.like(f"%{semester_name}%")
            ).first()
            if not semester:
                skipped += 1
                continue
            
            unit = db.query(Unit).filter(
                Unit.semester_id == semester.id,
                Unit.unit_number == int(unit_number) if unit_number else None
            ).first()
            if not unit:
                unit = db.query(Unit).filter(
                    Unit.semester_id == semester.id,
                    Unit.name == unit_name
                ).first()
            if not unit:
                skipped += 1
                continue
            
            exam_point = None
            if exam_point_title:
                exam_point = db.query(ExamPoint).filter(
                    ExamPoint.unit_id == unit.id,
                    ExamPoint.title == exam_point_title
                ).first()
            
            if question_type_name not in question_types_cache:
                qt = db.query(QuestionType).filter(QuestionType.name == question_type_name).first()
                if not qt:
                    qt = QuestionType(name=question_type_name)
                    db.add(qt)
                    db.commit()
                question_types_cache[question_type_name] = qt.id
            question_type_id = question_types_cache[question_type_name]
            
            if difficulty_name not in difficulties_cache:
                diff = db.query(Difficulty).filter(Difficulty.name == difficulty_name).first()
                if not diff:
                    diff = Difficulty(name=difficulty_name)
                    db.add(diff)
                    db.commit()
                difficulties_cache[difficulty_name] = diff.id
            difficulty_id = difficulties_cache[difficulty_name]
            
            import json
            question_json = None
            if json_content:
                try:
                    question_json = json.loads(json_content) if isinstance(json_content, str) else json_content
                except:
                    pass
            
            question = Question(
                version_id=grade.version_id,
                grade_id=grade.id,
                subject_id=subject.id,
                semester_id=semester.id,
                unit_id=unit.id,
                exam_point_id=exam_point.id if exam_point else None,
                question_type_id=question_type_id,
                difficulty_id=difficulty_id,
                content=stem_text or "",
                answer=answer_text or "",
                analysis=analysis_text or "",
                question_json=question_json,
                question_type=question_type_name,
                stem=stem_text
            )
            db.add(question)
            imported += 1
            
        except Exception as e:
            errors.append(f"第{row_idx}行: {str(e)}")
            skipped += 1
    
    db.commit()
    
    return {
        "message": "题库导入完成",
        "imported": imported,
        "skipped": skipped,
        "errors": errors[:10]
    }
